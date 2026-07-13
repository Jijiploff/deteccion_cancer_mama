import io
from datetime import datetime

import pandas as pd
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import inch, mm
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Image, Table, TableStyle,
    PageBreak,
)
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side


def generate_pdf(
    results: dict,
    cv_results: list,
    tuning_result: dict,
    stats_df: pd.DataFrame,
    eda_summary: dict,
    figures: dict,
    output_path: str,
):
    doc = SimpleDocTemplate(output_path, pagesize=A4,
                            rightMargin=2*cm, leftMargin=2*cm,
                            topMargin=2*cm, bottomMargin=2*cm)
    cm = mm * 10
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="CenterTitle", parent=styles["Title"],
                               alignment=TA_CENTER, fontSize=18, spaceAfter=20))
    styles.add(ParagraphStyle(name="SubTitle2", parent=styles["Heading2"],
                               fontSize=13, spaceAfter=8))

    elements = []

    elements.append(Paragraph("Reporte de Modelos - Detección de Cáncer de Mama",
                              styles["CenterTitle"]))
    elements.append(Paragraph(f"Generado: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
                              styles["Normal"]))
    elements.append(Spacer(1, 12))

    # 1. EDA Summary
    elements.append(Paragraph("1. Resumen del Dataset", styles["Heading1"]))
    eda_data = [
        ["Métrica", "Valor"],
        ["Calcificaciones", str(eda_summary.get("calcificaciones", "N/A"))],
        ["Masas", str(eda_summary.get("masas", "N/A"))],
        ["Casos Wisconsin", str(eda_summary.get("wisconsin", "N/A"))],
        ["Metadatos", str(eda_summary.get("metadatos", "N/A"))],
    ]
    t = Table(eda_data, colWidths=[200, 100])
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#8B5CF6")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 10),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F3E8FF")]),
    ]))
    elements.append(t)
    elements.append(Spacer(1, 12))

    if "pathology" in figures:
        img = Image(figures["pathology"], width=400, height=200)
        elements.append(img)
        elements.append(Spacer(1, 12))

    # 2. Training Results
    elements.append(Paragraph("2. Resultados de Entrenamiento", styles["Heading1"]))
    metrics_data = [["Modelo", "Accuracy", "Precision", "Recall", "F1", "AUC", "Tiempo (s)"]]
    for name, m in results.items():
        metrics_data.append([
            name,
            f"{m['accuracy']:.4f}",
            f"{m['precision']:.4f}",
            f"{m['recall']:.4f}",
            f"{m['f1']:.4f}",
            f"{m['auc']:.4f}",
            f"{m.get('training_time_s', 0):.3f}",
        ])
    t = Table(metrics_data, colWidths=[100, 70, 70, 70, 70, 70, 70])
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#8B5CF6")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F3E8FF")]),
    ]))
    elements.append(t)
    elements.append(Spacer(1, 12))

    best_model = max(results.items(), key=lambda x: x[1]["auc"])
    elements.append(Paragraph(
        f"<b>Mejor modelo:</b> {best_model[0]} (AUC = {best_model[1]['auc']:.4f})",
        styles["Normal"]
    ))
    elements.append(Spacer(1, 10))

    for name in results:
        if f"cm_{name}" in figures:
            elements.append(Paragraph(f"Matriz de Confusión - {name}", styles["SubTitle2"]))
            elements.append(Image(figures[f"cm_{name}"], width=280, height=240))
            elements.append(Spacer(1, 8))

    if "roc" in figures:
        elements.append(Paragraph("Curvas ROC", styles["SubTitle2"]))
        elements.append(Image(figures["roc"], width=400, height=300))
        elements.append(Spacer(1, 12))

    if "metrics_bar" in figures:
        elements.append(Paragraph("Comparación de Métricas", styles["SubTitle2"]))
        elements.append(Image(figures["metrics_bar"], width=400, height=260))
        elements.append(Spacer(1, 12))

    # 3. Cross-Validation
    if cv_results:
        elements.append(PageBreak())
        elements.append(Paragraph("3. Validación Cruzada", styles["Heading1"]))
        for cv in cv_results:
            elements.append(Paragraph(f"<b>{cv['model_name']}</b> - {cv['n_splits']}-Fold CV", styles["SubTitle2"]))
            cv_data = [["Fold", "Accuracy", "AUC"]]
            for r in cv["fold_results"]:
                cv_data.append([f"Fold {r['fold']}", f"{r['accuracy']:.4f}", f"{r['auc']:.4f}"])
            cv_data.append(["Media", f"{cv['mean_accuracy']:.4f}", f"{cv['mean_auc']:.4f}"])
            cv_data.append(["Desv. Est.", f"{cv['std_accuracy']:.4f}", f"{cv['std_auc']:.4f}"])
            t = Table(cv_data, colWidths=[80, 100, 100])
            t.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#8B5CF6")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("BACKGROUND", (0, -2), (-1, -1), colors.HexColor("#E5E7EB")),
            ]))
            elements.append(t)
            elements.append(Spacer(1, 8))
            elements.append(Paragraph(
                f"Accuracy: {cv['mean_accuracy']:.4f} ± {cv['std_accuracy']:.4f} | "
                f"AUC: {cv['mean_auc']:.4f} ± {cv['std_auc']:.4f}",
                styles["Normal"]
            ))
            elements.append(Spacer(1, 10))

            if f"cv_{cv['model_name']}" in figures:
                elements.append(Image(figures[f"cv_{cv['model_name']}"], width=380, height=220))
                elements.append(Spacer(1, 8))

    # 4. Hyperparameter Tuning
    if tuning_result:
        elements.append(PageBreak())
        elements.append(Paragraph("4. Hiperparámetros (Tuning)", styles["Heading1"]))
        elements.append(Paragraph(f"<b>Modelo:</b> {tuning_result['model_name']}", styles["Normal"]))
        elements.append(Paragraph(f"<b>Mejor AUC (CV):</b> {tuning_result['best_score']:.4f}", styles["Normal"]))
        elements.append(Paragraph(f"<b>Combinaciones evaluadas:</b> {tuning_result['total_combinations']}", styles["Normal"]))
        elements.append(Spacer(1, 8))

        hp_data = [["Parámetro", "Mejor Valor"]]
        for k, v in tuning_result["best_params"].items():
            hp_data.append([k, str(v)])
        t = Table(hp_data, colWidths=[150, 150])
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#8B5CF6")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ]))
        elements.append(t)
        elements.append(Spacer(1, 12))

    # 5. Statistical Tests
    if stats_df is not None and not stats_df.empty:
        elements.append(PageBreak())
        elements.append(Paragraph("5. Pruebas Estadísticas", styles["Heading1"]))
        stat_data = [["Modelo A", "Modelo B", "McNemar p", "Wilcoxon p", "¿Significativo?"]]
        for _, row in stats_df.iterrows():
            try:
                mcn_p = f"{float(row['McNemar p-valor']):.4f}"
            except (ValueError, TypeError):
                mcn_p = str(row.get("McNemar p-valor", "-"))
            try:
                wil_p = f"{float(row['Wilcoxon p-valor']):.4f}"
            except (ValueError, TypeError):
                wil_p = str(row.get("Wilcoxon p-valor", "-"))
            stat_data.append([
                row["Modelo A"], row["Modelo B"], mcn_p, wil_p,
                "Sí" if any("significativas" in str(row.get(c, "")) for c in
                            ["McNemar Interpretación", "Wilcoxon Interpretación"]) else "No",
            ])
        t = Table(stat_data, colWidths=[80, 80, 70, 70, 70])
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#8B5CF6")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 8),
            ("ALIGN", (0, 0), (-1, -1), "CENTER"),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ]))
        elements.append(t)
        elements.append(Spacer(1, 12))

    doc.build(elements)


def generate_word(
    results: dict,
    cv_results: list,
    tuning_result: dict,
    stats_df: pd.DataFrame,
    eda_summary: dict,
    figures: dict,
    output_path: str,
):
    doc = Document()

    doc.add_heading("Reporte de Modelos - Detección de Cáncer de Mama", level=0)
    doc.add_paragraph(f"Generado: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    # 1. EDA
    doc.add_heading("1. Resumen del Dataset", level=1)
    table = doc.add_table(rows=5, cols=2)
    table.style = "Light Shading Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    data = [
        ("Métrica", "Valor"),
        ("Calcificaciones", str(eda_summary.get("calcificaciones", "N/A"))),
        ("Masas", str(eda_summary.get("masas", "N/A"))),
        ("Casos Wisconsin", str(eda_summary.get("wisconsin", "N/A"))),
        ("Metadatos", str(eda_summary.get("metadatos", "N/A"))),
    ]
    for i, (k, v) in enumerate(data):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v

    if "pathology" in figures:
        doc.add_picture(figures["pathology"], width=Inches(5))

    # 2. Training
    doc.add_heading("2. Resultados de Entrenamiento", level=1)
    table = doc.add_table(rows=1 + len(results), cols=7)
    table.style = "Light Shading Accent 1"
    headers = ["Modelo", "Accuracy", "Precision", "Recall", "F1", "AUC", "Tiempo (s)"]
    for j, h in enumerate(headers):
        table.rows[0].cells[j].text = h
    for i, (name, m) in enumerate(results.items()):
        row = table.rows[i + 1]
        row.cells[0].text = name
        row.cells[1].text = f"{m['accuracy']:.4f}"
        row.cells[2].text = f"{m['precision']:.4f}"
        row.cells[3].text = f"{m['recall']:.4f}"
        row.cells[4].text = f"{m['f1']:.4f}"
        row.cells[5].text = f"{m['auc']:.4f}"
        row.cells[6].text = f"{m.get('training_time_s', 0):.3f}"

    for name in results:
        if f"cm_{name}" in figures:
            doc.add_heading(f"Matriz de Confusión - {name}", level=2)
            doc.add_picture(figures[f"cm_{name}"], width=Inches(4))

    if "metrics_bar" in figures:
        doc.add_heading("Comparación de Métricas", level=2)
        doc.add_picture(figures["metrics_bar"], width=Inches(5))

    # 3. CV
    if cv_results:
        doc.add_heading("3. Validación Cruzada", level=1)
        for cv in cv_results:
            doc.add_heading(cv["model_name"], level=2)
            table = doc.add_table(rows=3 + cv["n_splits"], cols=3)
            table.style = "Light Shading Accent 1"
            table.rows[0].cells[0].text = "Fold"
            table.rows[0].cells[1].text = "Accuracy"
            table.rows[0].cells[2].text = "AUC"
            for i, r in enumerate(cv["fold_results"]):
                table.rows[i + 1].cells[0].text = f"Fold {r['fold']}"
                table.rows[i + 1].cells[1].text = f"{r['accuracy']:.4f}"
                table.rows[i + 1].cells[2].text = f"{r['auc']:.4f}"
            table.rows[cv["n_splits"] + 1].cells[0].text = "Media"
            table.rows[cv["n_splits"] + 1].cells[1].text = f"{cv['mean_accuracy']:.4f}"
            table.rows[cv["n_splits"] + 1].cells[2].text = f"{cv['mean_auc']:.4f}"
            table.rows[cv["n_splits"] + 2].cells[0].text = "Desv. Est."
            table.rows[cv["n_splits"] + 2].cells[1].text = f"{cv['std_accuracy']:.4f}"
            table.rows[cv["n_splits"] + 2].cells[2].text = f"{cv['std_auc']:.4f}"

    # 4. Hyperparameter Tuning
    if tuning_result:
        doc.add_heading("4. Hiperparámetros", level=1)
        doc.add_paragraph(f"Mejor AUC (CV): {tuning_result['best_score']:.4f}")
        table = doc.add_table(rows=1 + len(tuning_result["best_params"]), cols=2)
        table.style = "Light Shading Accent 1"
        table.rows[0].cells[0].text = "Parámetro"
        table.rows[0].cells[1].text = "Mejor Valor"
        for i, (k, v) in enumerate(tuning_result["best_params"].items()):
            table.rows[i + 1].cells[0].text = k
            table.rows[i + 1].cells[1].text = str(v)

    # 5. Statistical Tests
    if stats_df is not None and not stats_df.empty:
        doc.add_heading("5. Pruebas Estadísticas", level=1)
        table = doc.add_table(rows=1 + len(stats_df), cols=5)
        table.style = "Light Shading Accent 1"
        for j, h in enumerate(["Modelo A", "Modelo B", "McNemar p", "Wilcoxon p", "¿Significativo?"]):
            table.rows[0].cells[j].text = h
        for i, (_, row) in enumerate(stats_df.iterrows()):
            table.rows[i + 1].cells[0].text = row["Modelo A"]
            table.rows[i + 1].cells[1].text = row["Modelo B"]
            try:
                table.rows[i + 1].cells[2].text = f"{float(row['McNemar p-valor']):.4f}"
            except (ValueError, TypeError):
                table.rows[i + 1].cells[2].text = str(row.get("McNemar p-valor", "-"))
            try:
                table.rows[i + 1].cells[3].text = f"{float(row['Wilcoxon p-valor']):.4f}"
            except (ValueError, TypeError):
                table.rows[i + 1].cells[3].text = str(row.get("Wilcoxon p-valor", "-"))
            table.rows[i + 1].cells[4].text = "Sí" if "significativas" in str(
                row.get("McNemar Interpretación", "")
            ) or "significativas" in str(row.get("Wilcoxon Interpretación", "")) else "No"

    doc.save(output_path)


def generate_excel(
    results: dict,
    cv_results: list,
    tuning_result: dict,
    stats_df: pd.DataFrame,
    eda_summary: dict,
    output_path: str,
):
    wb = Workbook()
    header_font = Font(bold=True, color="FFFFFF", size=11)
    header_fill = PatternFill(start_color="8B5CF6", end_color="8B5CF6", fill_type="solid")
    green_fill = PatternFill(start_color="D4EDDA", end_color="D4EDDA", fill_type="solid")
    thin_border = Border(
        left=Side(style="thin"), right=Side(style="thin"),
        top=Side(style="thin"), bottom=Side(style="thin"),
    )

    # Sheet 1: Metrics
    ws1 = wb.active
    ws1.title = "Métricas"
    headers = ["Modelo", "Accuracy", "Precision", "Recall", "F1", "AUC", "Tiempo (s)"]
    for j, h in enumerate(headers, 1):
        cell = ws1.cell(row=1, column=j, value=h)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center")
    for i, (name, m) in enumerate(results.items(), 2):
        ws1.cell(row=i, column=1, value=name)
        ws1.cell(row=i, column=2, value=round(m["accuracy"], 4))
        ws1.cell(row=i, column=3, value=round(m["precision"], 4))
        ws1.cell(row=i, column=4, value=round(m["recall"], 4))
        ws1.cell(row=i, column=5, value=round(m["f1"], 4))
        ws1.cell(row=i, column=6, value=round(m["auc"], 4))
        ws1.cell(row=i, column=7, value=m.get("training_time_s", 0))
        for j in range(1, 8):
            ws1.cell(row=i, column=j).border = thin_border
            ws1.cell(row=i, column=j).alignment = Alignment(horizontal="center")
    ws1.column_dimensions["A"].width = 18
    for col in range(2, 8):
        ws1.column_dimensions[chr(64 + col)].width = 14

    best_row = 2 + list(results.values()).index(max(results.values(), key=lambda x: x["auc"]))
    for j in range(1, 8):
        ws1.cell(row=best_row, column=j).fill = green_fill

    # Sheet 2: Cross-Validation
    if cv_results:
        ws2 = wb.create_sheet("Validación Cruzada")
        row_idx = 1
        for cv in cv_results:
            ws2.cell(row=row_idx, column=1, value=f"{cv['model_name']} - {cv['n_splits']}-Fold CV")
            ws2.cell(row=row_idx, column=1).font = Font(bold=True, size=12)
            row_idx += 1
            for j, h in enumerate(["Fold", "Accuracy", "AUC"], 1):
                cell = ws2.cell(row=row_idx, column=j, value=h)
                cell.font = header_font
                cell.fill = header_fill
                cell.alignment = Alignment(horizontal="center")
            row_idx += 1
            for r in cv["fold_results"]:
                ws2.cell(row=row_idx, column=1, value=f"Fold {r['fold']}")
                ws2.cell(row=row_idx, column=2, value=round(r["accuracy"], 4))
                ws2.cell(row=row_idx, column=3, value=round(r["auc"], 4))
                for j in range(1, 4):
                    ws2.cell(row=row_idx, column=j).border = thin_border
                row_idx += 1
            ws2.cell(row=row_idx, column=1, value="Media").font = Font(bold=True)
            ws2.cell(row=row_idx, column=2, value=round(cv["mean_accuracy"], 4))
            ws2.cell(row=row_idx, column=3, value=round(cv["mean_auc"], 4))
            for j in range(1, 4):
                ws2.cell(row=row_idx, column=j).border = thin_border
            row_idx += 1
            ws2.cell(row=row_idx, column=1, value="Desv. Est.").font = Font(bold=True)
            ws2.cell(row=row_idx, column=2, value=round(cv["std_accuracy"], 4))
            ws2.cell(row=row_idx, column=3, value=round(cv["std_auc"], 4))
            for j in range(1, 4):
                ws2.cell(row=row_idx, column=j).border = thin_border
            row_idx += 2

    # Sheet 3: Hyperparameters
    if tuning_result:
        ws3 = wb.create_sheet("Hiperparámetros")
        ws3.cell(row=1, column=1, value="Parámetro").font = header_font
        ws3.cell(row=1, column=1).fill = header_fill
        ws3.cell(row=1, column=2, value="Mejor Valor").font = header_font
        ws3.cell(row=1, column=2).fill = header_fill
        for i, (k, v) in enumerate(tuning_result["best_params"].items(), 2):
            ws3.cell(row=i, column=1, value=k)
            ws3.cell(row=i, column=2, value=str(v))
        ws3.column_dimensions["A"].width = 20
        ws3.column_dimensions["B"].width = 20

    # Sheet 4: Statistical Tests
    if stats_df is not None and not stats_df.empty:
        ws4 = wb.create_sheet("Pruebas Estadísticas")
        for j, h in enumerate(["Modelo A", "Modelo B", "McNemar p", "Wilcoxon p", "¿Significativo?"], 1):
            cell = ws4.cell(row=1, column=j, value=h)
            cell.font = header_font
            cell.fill = header_fill
        for i, (_, row) in enumerate(stats_df.iterrows(), 2):
            ws4.cell(row=i, column=1, value=row["Modelo A"])
            ws4.cell(row=i, column=2, value=row["Modelo B"])
            try:
                ws4.cell(row=i, column=3, value=round(float(row["McNemar p-valor"]), 4))
            except (ValueError, TypeError):
                ws4.cell(row=i, column=3, value=str(row.get("McNemar p-valor", "-")))
            try:
                ws4.cell(row=i, column=4, value=round(float(row["Wilcoxon p-valor"]), 4))
            except (ValueError, TypeError):
                ws4.cell(row=i, column=4, value=str(row.get("Wilcoxon p-valor", "-")))
            ws4.cell(row=i, column=5, value="Sí" if "significativas" in str(
                row.get("McNemar Interpretación", "")
            ) else "No")
            for j in range(1, 6):
                ws4.cell(row=i, column=j).border = thin_border
                ws4.cell(row=i, column=j).alignment = Alignment(horizontal="center")

    wb.save(output_path)
